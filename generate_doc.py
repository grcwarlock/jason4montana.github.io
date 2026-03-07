"""Generate a Word document from the jason4montana.github.io website copy."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- Styles helpers ---
def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    return p

def heading3(text):
    return doc.add_heading(text, level=3)

def label(text):
    """Small caps section label."""
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    return p

def body(text):
    return doc.add_paragraph(text)

def pull_quote(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Quote'] if 'Quote' in [s.name for s in doc.styles] else doc.styles['Normal']
    for run in p.runs:
        run.italic = True
    return p

def page_break():
    doc.add_page_break()

def divider():
    doc.add_paragraph("─" * 60)

# ================================================================
# TITLE PAGE
# ================================================================
title = doc.add_heading("Jason for Montana HD-37", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph("Website Copy — Draft for Review")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub.runs:
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph("")
contact = doc.add_paragraph("jason@jason4montana.com  ·  (406) 366-0318  ·  Lewistown, MT  ·  2026")
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in contact.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

page_break()

# ================================================================
# INDEX.HTML — HOME PAGE
# ================================================================
label("Page")
heading1("Home Page  (index.html)")

label("Hero")
heading2("No party. Just solutions.")
body("Tagline / eyebrow: Montana House District 37  ·  Independent Candidate  ·  2026")
body("Candidate name displayed: Jason Wilson")

doc.add_paragraph("")
label("The Story — Intro Section")
body(
    "Growing up a military kid means a new base every three years, a new school, new faces to figure out. "
    "My grandmother's house on West Boulevard in Lewistown was the one address that never changed, and summers here "
    "were what home actually felt like. A lot of world came after those childhood, teenage, and young adult summers. "
    "Securing this country as an intelligence officer at the Defense Intelligence Agency, and later at the Pentagon "
    "during the Global War on Terror, took me about as far from West Boulevard as a person can get. Work in "
    "cybersecurity and data privacy followed, spanning public sector and private sector both, and across all of it "
    "the same skill kept showing up: knowing how to find what is actually true underneath the version someone is "
    "selling you. Bringing that home to the only place that was ever really mine felt less like a decision and more "
    "like the obvious next thing. Running as an independent means my obligation runs in one direction only, to the "
    "residents of District 37."
)

doc.add_paragraph("")
label("Card — Who I Am")
heading3("Intelligence officer. Four generations of Central Montana. No party obligations.")
body(
    "From the DIA and the Pentagon to the only hometown I ever had. That combination gives me a particular way of "
    "seeing what is actually heading toward this region right now."
)

label("Card — The Platform")
heading3("Five issues. Real positions. No consultant language.")
body(
    "Data centers, housing affordability, Montana producers, rural economic development, and working Montanans. "
    "Every position on this site is something I will say to your face at a kitchen table."
)

label("Accent Band")
body("Local. No party money. One district. One obligation.")

page_break()

# ================================================================
# ABOUT.HTML
# ================================================================
label("Page")
heading1("About Page  (about.html)")

heading2("Four generations in. The world got my time. This place gets the rest of it.")

label("Lede")
body(
    "Growing up a military kid means a new base every three years, a new school, new faces to figure out. "
    "My grandmother's house on West Boulevard was the one address that never changed, and summers here were what "
    "home actually felt like: Fourth of July fireworks off the driveway, days spent in the Judith Mountains, the "
    "particular feeling of a place that belongs to you even when you cannot quite claim to live there yet. Four "
    "generations of this family built something in Central Montana, and I spent my childhood counting down to when "
    "I could come back."
)

doc.add_paragraph("")
heading2("What I did out there.")
body(
    "A lot of world came after those childhood, teenage, and young adult summers. Securing this country as an "
    "intelligence officer at the Defense Intelligence Agency, and later at the Pentagon during the Global War on "
    "Terror, took me about as far from West Boulevard as a person can get. Work in cybersecurity and data privacy "
    "followed, spanning public sector and private sector both, and across all of it the same skill kept showing up: "
    "knowing how to find what is actually true underneath the version someone is selling you."
)
body(
    "That skill is not abstract. It is the difference between reading a data center incentive bill and reading the "
    "summary someone prepared about it, between understanding what a water consumption commitment actually requires "
    "and what it merely suggests, between knowing what a legislative deal does and what its press release claims it "
    "does. Bringing it home to represent Central Montana in Helena is the most direct use I can imagine for "
    "everything those years taught me."
)

pull_quote(
    "The decisions that will shape what this region looks like in twenty years are being made right now, mostly by "
    "people who do not live here and do not plan to. Having someone in that room who knows how to read what is "
    "actually in the agreement, not just the version designed for public consumption, is the whole point of running."
)

doc.add_paragraph("")
heading2("The roots that brought me back.")
body(
    "My great-grandfather came to Lewistown in 1911 and became the town's first osteopath. My grandmother was born "
    "here in 1920. My grandfather worked as a mining engineer for the Anaconda Company, which is how my mother came "
    "into the world in Butte; Central Montana families have always followed the work, and they have always come back. "
    "Running AI and cybersecurity literacy workshops for neighbors here started well before any campaign did, because "
    "people had real questions about the technology being deployed into their community and deserved real answers. "
    "Organizing the community response to the zoning variance last year grew from the same instinct: residents "
    "engaged in good faith, made legally grounded arguments, and were overruled. Learning exactly why that happened, "
    "and what tools exist at the state level to change it, is the kind of groundwork that does not make the paper "
    "and does, however, translate directly into useful representation."
)
body(
    "Watching outside interests make consequential decisions about Central Montana from a considerable distance, "
    "without meaningful input from the people who actually live here, is personal rather than abstract. That is a "
    "feature, not a problem, in a representative."
)

doc.add_paragraph("")
heading2("Why independent, and why this race.")
body(
    "Both party candidates will arrive in Helena carrying obligations that predate anything District 37 asked of "
    "them: caucus relationships, state party leadership expectations, and donor networks operating well above the "
    "district level create real and ongoing pressure that runs in one direction. Running as an independent removes "
    "all of it. The accountability in this campaign runs horizontally, to the residents of this district, full stop."
)
body(
    "Four generations of this family lived and worked and built things here. Representing Central Montana in the "
    "state house is not a career move or a pivot. It is the logical next thing, and it is long past time someone "
    "did it without a party telling them how."
)

page_break()

# ================================================================
# SOLUTIONS.HTML
# ================================================================
label("Page")
heading1("Solutions Page  (solutions.html)")

heading2("Five issues. Straight answers.")
body(
    "Every position on this site is something I will say to your face at a kitchen table. No consultant language, "
    "no focus-grouped equivocation. These are the things that matter in District 37 right now."
)

doc.add_paragraph("")

# Issue 01
label("Issue 01")
heading2("Working Montanans — Policy should make daily life more manageable, not more complicated.")
body(
    "Wages have not kept pace with the cost of living in most of Montana, and the legislature has consistently "
    "prioritized corporate tax environments over the conditions that let working families stay stable. Healthcare "
    "access in rural communities, childcare availability, and basic housing affordability all deserve direct "
    "legislative attention. The market-will-sort-it-out answer has been the standing answer in this district for "
    "a long time; the results are visible to anyone paying attention."
)
body(
    "Working Montanans are not asking for complicated policy architecture. They are asking whether their wages "
    "cover their rent, whether there is a doctor within a reasonable drive, and whether childcare costs less than "
    "a second job would bring in. Solvable problems, all of them, provided the legislature treats them as "
    "priorities rather than talking points deployed before primaries and forgotten about afterward."
)
pull_quote(
    "Pragmatic representation means asking what works, what can pass, and what actually reaches the people it is "
    "intended to help. Ideological purity on either side of these questions has not produced results in this "
    "district. Something more direct is long overdue."
)
body(
    "The nurses at the hospital, the teachers in the schools, the mechanics, the tradespeople, the people running "
    "businesses on Main Street — these are the people this campaign is built for, not an abstract version of their "
    "interests. Showing up for them specifically is the job."
)

doc.add_paragraph("")

# Issue 02
label("Issue 02")
heading2("Montana Producers — Montana raises it. Someone else is making the money on it.")
body(
    "A steer raised on a Central Montana ranch, grass-fed and well-managed, the product of a family that has "
    "worked the same ground for three or four generations, goes to a processing facility in another state, gets "
    "packaged under a brand owned by a corporation headquartered somewhere else, and lands in a grocery store with "
    "nothing on the label indicating it came from Montana. The rancher gets a commodity price. Someone else builds "
    "the margin. The value that this state's land, water, and generations of accumulated knowledge created flows "
    "out of Montana as reliably as the seasons turn."
)
body(
    "Direct-to-consumer sales change that math in a fundamental way. A family ranch that ships beef directly to "
    "subscribers across the country, with their name alongside the Made in Montana mark, keeps the margin here. "
    "Customers who care about where their food comes from will pay a meaningful premium for that direct connection; "
    "the relationship builds over time into something more durable than a commodity contract and far more resilient "
    "than commodity price swings. This is not a niche play for boutique buyers. It is a structural shift in how "
    "agricultural value flows, already happening in states that have made it legislatively straightforward."
)
pull_quote(
    "Montana produces things people will pay a premium for, specifically because they come from here and because "
    "knowing who made them matters. Unlocking that value means removing the regulatory barriers that make "
    "direct-to-consumer shipping difficult for small and mid-size producers, strengthening the Made in Montana "
    "label as a genuine quality signal, and putting real legislative support behind keeping multi-generational "
    "farm and ranch families on their land."
)
body(
    "Supporting the Made in Montana program means more than a logo. Building traceability into the supply chain, "
    "supporting small-scale USDA-inspected processing infrastructure in rural Montana, and making it possible for "
    "a family operation to reach customers anywhere in the country without surrendering most of the value to a "
    "middleman all belong in the same conversation. So does succession planning, because losing a "
    "multi-generational ranch is not a market transaction. The land endures; the families who know it and love it "
    "are considerably harder to replace."
)

doc.add_paragraph("")

# Issue 03
label("Issue 03")
heading2("Housing Affordability — Working people should be able to afford to live where they work.")
body(
    "Housing costs across Montana have climbed sharply over the past several years, driven by outside investment, "
    "limited inventory, and a construction market that responds to demand at the higher end of the price range long "
    "before it touches the working end. For people who grew up here, work here, and want to stay here, the gap "
    "between wages and what a decent house actually costs has become a real and practical barrier. Not a talking "
    "point; a lived condition that shows up in who leaves and who stays."
)
body(
    "Rural communities face a version of this problem distinct from what is playing out in Missoula or Bozeman. "
    "Stagnant wage growth relative to a housing stock that has not meaningfully expanded in decades is the driver "
    "here, and modest-income families, teachers, medical workers, and tradespeople are getting priced out quietly, "
    "one lease renewal at a time. The people who keep this community functioning deserve to afford a home in it."
)
pull_quote(
    "Workforce housing investment, zoning reform that actually enables construction of modest homes, and policies "
    "that prioritize long-term residents over short-term investment returns are the real levers available in Helena. "
    "Using them requires treating housing affordability as an economic necessity rather than a social program."
)
body(
    "Local communities should still shape how they grow, and residents deserve genuine authority over what gets "
    "built in their neighborhoods. The state's role is to set conditions that make affordable housing easier to "
    "build and harder to speculate away. Right now it largely does neither, and that is a policy choice, not an "
    "inevitability."
)

doc.add_paragraph("")

# Issue 04
label("Issue 04")
heading2("Rural Economic Development — Real investment looks different from the press release version.")
body(
    "Outside capital that arrives with fanfare, captures available tax incentives, repatriates profit to distant "
    "shareholders, and leaves when the incentive structure changes is not economic development for Central Montana. "
    "It is a lease on the appearance of progress, paid for by the public. Real investment stays, hires locally, "
    "pays into the tax base, and builds something that functions after the ribbon-cutting photo fades from the "
    "local paper."
)
body(
    "Broadband access remains foundational infrastructure for rural Montana and remains incomplete throughout much "
    "of this district. Small business support, agricultural infrastructure, and keeping young people here rather "
    "than watching them leave for Billings are the actual economic foundations of a community like Lewistown. These "
    "receive considerably less legislative attention than large headline projects, and the slow drift of population "
    "and commerce that any long-time resident can name without looking at a census table is the result."
)
pull_quote(
    "Economic development policy that works for Central Montana looks different from policy written for the I-90 "
    "corridor. A representative from this district should be making that case in Helena rather than deferring to "
    "frameworks designed for someone else's community."
)
body(
    "Prioritizing local ownership structures, enforceable local hiring requirements in public contracts, and "
    "broadband investment over data center tax breaks reflects a simple judgment about whose economy is actually "
    "being developed. The answer should be the people who were already here, and a legislature that means it "
    "should be able to show its work."
)

doc.add_paragraph("")

# Issue 05
label("Issue 05")
heading2("Data Centers & Tech Infrastructure — Montana's infrastructure should serve Montanans first.")
body(
    "Large-scale data center developments are moving through the legislature with minimal public scrutiny and "
    "substantial public subsidy. Facilities of this size consume enormous amounts of water and electricity, "
    "generate relatively few local jobs for the footprint they require, and lock communities into long-term "
    "infrastructure obligations that outlast the political moment that created them. Reading the actual bills, "
    "rather than the summaries prepared by the people who want them to pass, tells a meaningfully different story "
    "than the press releases do."
)
body(
    "Central Montana sits in a region where water is not an abundant resource and where rural power reliability "
    "matters to everyone who lives here year-round. Committing meaningful portions of that capacity to serve "
    "outside interests is a decision that demands an honest accounting before any vote. Tax incentive structures "
    "being written right now carry limited clawback provisions; a project that fails to deliver on its promised "
    "employment figures faces minimal financial consequence, and that risk lands on the public, not the developer."
)
pull_quote(
    "Any project of this scale should require full public disclosure before approval, a genuine accounting of "
    "water and grid impacts, and local hiring provisions with real enforcement. \"Economic development\" is not a "
    "blank check, and proximity to a ribbon-cutting photo does not make something good for the people who were "
    "already here."
)
body(
    "Spending years as an intelligence officer, where the gap between what an agreement says on paper and what it "
    "actually commits to is a professional concern, gives me a particular orientation toward documents like these. "
    "Bringing that reading to the legislative process is a specific thing this candidacy offers that neither party "
    "candidate can match."
)

page_break()

# ================================================================
# NEWS ARTICLES
# ================================================================
label("Page")
heading1("News Articles  (news/)")

# Article 1
label("Article — Agriculture  ·  February 2026")
heading2("Montana's Ranchers Are Feeding the Country. Someone Else Is Cashing the Check.")
label("Lede")
body(
    "Keeping the money in Montana means selling the steak, not just the steer. The Made in Montana label and "
    "direct-to-consumer shipping are two of the most underused tools available for changing that equation, and "
    "both need a legislature that actually shows up for producers."
)
body(
    "Consider what happens to a steer raised on a Central Montana ranch, grass-fed and well-managed, the product "
    "of a family that has worked the same ground for three or four generations. That animal goes to a processing "
    "facility in another state, gets packaged under a brand owned by a corporation headquartered somewhere else, "
    "and lands in a grocery store with nothing on the label indicating it came from Montana. The rancher gets a "
    "commodity price. Someone else builds the margin. The value that this state's land, water, and generations of "
    "accumulated knowledge created flows out of Montana as reliably as a spring runoff, with nothing coming back "
    "to match it."
)
body(
    "Run the same operation through a direct-to-consumer model and the picture changes completely. A family ranch "
    "that ships beef directly to subscribers across the country, with their name alongside the Made in Montana "
    "mark, keeps the margin here. Customers who care about where their food comes from will pay a meaningful "
    "premium for that direct connection; the relationship builds over time into something more durable than a "
    "commodity contract and considerably more resilient than commodity price swings. Calling this a niche market "
    "misses what it actually is: a structural shift in how agricultural value flows, already happening in states "
    "that have made it legislatively straightforward."
)
body(
    "The barriers are not mysterious. Direct-to-consumer meat shipping involves a tangle of USDA inspection "
    "requirements, state licensing, and carrier logistics that large operations can navigate and that genuinely "
    "burdens a family ranch. Strengthening the Made in Montana program means building a traceability "
    "infrastructure that small producers can plug into rather than build from scratch, one that lets a buyer "
    "anywhere in the country know exactly which operation their food came from. Support for small-scale "
    "USDA-inspected processing facilities in rural Montana removes one of the biggest bottlenecks in the chain. "
    "None of this requires complicated policy; it requires a legislature that treats agricultural producers as a "
    "constituency worth investing in, not just a backdrop for campaign photos taken once every two years."
)
body(
    "The multi-generational dimension matters at least as much as the economic argument. A ranch that has been in "
    "a family since homesteading is not simply a business, and losing one to a combination of cost pressure, "
    "estate complications, and lack of succession infrastructure is a permanent change to what Montana is. The "
    "land endures; the families who know it and love it are considerably harder to replace."
)

doc.add_paragraph("")

# Article 2
label("Article — Infrastructure  ·  February 2026")
heading2("What the Data Center Bills Actually Say")
label("Lede")
body(
    "The legislation moving through Helena looks like economic development on the surface. Reading the full "
    "statutory text is a considerably different experience, and the difference matters for Central Montana."
)
body(
    "Two bills working through the current legislative session, both related to data center development "
    "incentives, have received relatively little public attention given what they actually authorize. Testimony "
    "in support came primarily from development interests. The public comment window moved on a timeline that "
    "made substantive community organizing difficult, which is not unusual for infrastructure legislation and is "
    "worth paying close attention to regardless."
)
body(
    "Reading the full text rather than the one-page summary circulated by supporters reveals several provisions "
    "that deserve scrutiny from any legislator representing a rural district with constrained water and power "
    "resources. Tax abatement structures in both bills carry limited clawback provisions; developments that "
    "underperform on promised employment figures face minimal financial consequence, and that risk lands on the "
    "public, not the developer. Water consumption estimates in supporting testimony rely on operational "
    "projections rather than contractual commitments, and for a region where agricultural water rights are "
    "already contested, that distinction is material."
)
body(
    "Spending years as an intelligence officer at the DIA and Pentagon, where the gap between what an agreement "
    "says and what it actually commits to is a professional concern, gives me a particular orientation toward "
    "documents like these. Both party candidates have structural incentives to align with their leadership's "
    "position on economic development legislation. An independent member has one incentive: whether the bill "
    "actually serves the people who sent them to Helena. That is the lens I would bring to these votes, and it "
    "is one the current conversation is short of."
)

doc.add_paragraph("")

# Article 3
label("Article — Campaign  ·  January 2026")
heading2("Why I'm Running Independent in a Three-Way Race")
label("Lede")
body(
    "Two party candidates, one district, and a straightforward question about who a representative actually works "
    "for once the election is over."
)
body(
    "Growing up a military kid means a new base every three years, a new school, a new set of faces to figure "
    "out. My grandmother's house on West Boulevard in Lewistown was the one address that never changed, and "
    "summers here were what home actually felt like. Four generations of this family put down roots in Central "
    "Montana, and I spent my childhood counting the days until I got to come back. A lot of world came after "
    "those summers, including years as an intelligence officer at the Defense Intelligence Agency and later at "
    "the Pentagon during the Global War on Terror, and work in cybersecurity and data privacy spanning public "
    "sector and private sector both. Bringing all of it back to the only place that was ever really mine felt "
    "less like a decision and more like the obvious next thing."
)
body(
    "The skill that running intelligence operations builds, specifically, is reading past the version of events "
    "someone needs you to believe and finding what is actually there. Bills get written with specific "
    "beneficiaries in mind. Regulatory frameworks contain carve-outs invisible in the summary version. "
    "Testimony coordinated well in advance shapes the legislative record in ways that public comment rarely "
    "corrects. That capacity, brought to the work of representing District 37 in Helena, is the specific thing "
    "this candidacy offers."
)
body(
    "Both party candidates will arrive carrying obligations that predate anything this district asked of them: "
    "caucus relationships, state party expectations, and donor networks operating well above the district level "
    "create real and ongoing pressure that runs in one direction. Running as an independent removes all of it. "
    "The accountability here runs to the residents of District 37, full stop. No party. Just solutions. That is "
    "the whole argument, and it is the one I intend to make all the way to November."
)

# Save
output_path = "/home/user/jason4montana.github.io/jason-for-montana-webcopy.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
